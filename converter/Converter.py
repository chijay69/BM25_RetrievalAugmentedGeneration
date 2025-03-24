import logging
import docx
import pypandoc
from config import CONVERTER_LOG_PATH


class Converter:
    """
    Envelops helper methods to convert various text formats to Markdown text format
    """

    def __init__(self, file_input_path: str = None):
        """
        Initialize the class.
        """
        self.txt_file = None
        self.input_path: str = file_input_path
        logging.basicConfig(level=logging.INFO,  # Set the minimum logging level
                            format='%(asctime)s - %(levelname)s - %(message)s',  # Define the log message format
                            filename=CONVERTER_LOG_PATH,  # Optional: Log to a file
                            filemode='a')  # Optional: Append to the log file (default), or 'w' to overwrite

        self.log = logging.getLogger("Converter")


    def check_file_path(self):
        if self.input_path is None:
            raise FileExistsError("Input path cannot be None.")

    # Function to convert docx to markdown
    @classmethod
    def convert_docx_to_markdown(cls, input_path: str) -> str:
        """
        This function converts a docx file to a markdown file.
        :param input_path: The file path of the file to be read and converted.
        :return: A Markdown format of the file read.
        """
        doc = docx.Document(input_path)
        md_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return md_text

    @classmethod
    # Function to convert DOC to Markdown using pypandoc
    def convert_doc_to_markdown(cls, input_path1: str) -> str:
        """
        Converts a document to markdown using the pypandoc module.
        :param input_path1: The path to the doc file.
        :return: The markdown file.
        """
        TEXT_EXT_TYPE = "md"
        return pypandoc.convert_file(input_path1, TEXT_EXT_TYPE)

    @classmethod
    def convert_docx_to_markdown(cls, input_path) -> str:
        """
        This function converts a docx file to a markdown file.
        :param input_path: The file path of the file to be read and converted.
        :return: A Markdown format of the file read.
        """
        doc = docx.Document(input_path)
        md_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return md_text

    @classmethod
    # Placeholder function for text to markdown
    def convert_txt_to_markdown(cls, input_path2: str) -> str:
        """
        Converts a text string to a markdown formatted document.
        :param input_path2: The path to the doc file.
        :return: The markdown file.
        """
        TEXT_EXT_TYPE = "md"

        with open(input_path2, 'r') as file:
            txt_file = file.read()

        return pypandoc.convert_text(txt_file, TEXT_EXT_TYPE, 'str')

    # Function to convert DOC to Markdown using pypandoc
    def convert_doc_to_markdown(self) -> str|None:
        """
        Converts a document to markdown using the pypandoc module.
        :param self: The path to the doc file.
        :return: The markdown file.
        """
        TEXT_EXT_TYPE = "md"
        try:
            self.check_file_path()
        except Exception as e:
            self.log.exception("Set your input path to a valid file path", exc_info=e)
            return None
        return pypandoc.convert_file(self.input_path, TEXT_EXT_TYPE)

    def convert_docx_to_markdown(self) -> str | None:
        """
        This function converts a docx file to a markdown file.
        :param input_path: The file path of the file to be read and converted.
        :return: A Markdown format of the file read.
        """
        try:
            self.check_file_path()
        except Exception as e:
            self.log.exception("Set your input path to a valid file path", exc_info=e)
            return None
        doc = docx.Document(self.input_path)
        md_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return md_text

    # Placeholder function for text to markdown
    def convert_txt_to_markdown(self) -> str|None:
        """
        Converts a text string to a markdown formatted document.
        :param self: The path to the doc file.
        :return: The markdown file.
        """
        TEXT_EXT_TYPE = "md"
        try:
            self.check_file_path()
        except Exception as e:
            self.log.exception("Set your input path to a valid file path", exc_info=e)
            return None
        with open(self.input_path, 'r') as file:
            self.txt_file = file.read()

        return pypandoc.convert_text(self.txt_file, TEXT_EXT_TYPE, 'str')

