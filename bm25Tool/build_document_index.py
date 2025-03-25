"""A python file with function build document index"""
import logging
import os
import re
from typing import List, Dict, Set, Any, Tuple, Optional

import docx
import nltk
import pymupdf4llm

from bm25Tool.setup_logger import setup_logger
from config_reader import get_output_dir, get_base_directory, get_data_dir, get_chunk_size
from converter.Document import Document

BASE_DIR = get_base_directory()

CONFIG_PATH = os.path.join(BASE_DIR, "config.ini")
DATA_PATH = os.path.join(BASE_DIR, get_data_dir(CONFIG_PATH))
OUTPUT_PATH = os.path.join(BASE_DIR, get_output_dir(CONFIG_PATH))
LOG_PATH = os.path.join(BASE_DIR, 'logs/build_document.log')

FILE_TYPES: Set[str] = {".pdf", ".docx"}
chunk_size = get_chunk_size(CONFIG_PATH)

logger: logging.Logger = setup_logger(__file__)


def _convert_pdf_to_markdown(input_path: str, output_path: str) -> None:
    """Asynchronously converts PDF to Markdown using pymupdf4llm."""
    try:
        markdown_text: str = pymupdf4llm.to_markdown(input_path)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(markdown_text)
        logger.info(f"Successfully converted PDF: {input_path} to {output_path}")
    except FileNotFoundError as e:
        logger.error(f"PDF file not found: {input_path}. Error: {e}")
    except Exception:
        logger.exception(f"An unexpected error occurred during PDF conversion: {input_path}")

def _convert_docx_to_markdown(input_path: str, output_path: str) -> None:
    """Asynchronously converts DOCX to Markdown."""
    try:
        doc = docx.Document(input_path)
        markdown_text = ""
        for paragraph in doc.paragraphs:
            markdown_text += paragraph.text + "\n"
        with open(output_path, "w", encoding="utf-8") as f: #Include encoding
            f.write(markdown_text)
        logger.info(f"Successfully converted DOCX: {input_path} to {output_path}")
    except FileNotFoundError as e:
        logger.error(f"DOCX file not found: {input_path}. Error: {e}")
    except Exception:
        logger.exception(f"An unexpected error occurred during DOCX conversion: {input_path}")

def convert_files_to_markdown(
    input_dir: str, output_dir: str, file_extensions: Set[str]
) -> Dict[str, str]:
    """
    Asynchronously converts files with specified extensions to Markdown.

    Args:
        input_dir: Path to the input directory.
        output_dir: Path to the output directory.
        file_extensions: Set of file extensions to convert (e.g., {'.pdf', '.docx'}).

    Returns:
        A dictionary mapping input file paths to output file paths.

    Raises:
        FileNotFoundError: If the input directory does not exist.
        OSError: If there are issues creating the output directory or accessing files.
        Exception: For any other unexpected errors during conversion.
    """
    try:
        if not os.path.exists(input_dir):
            raise FileNotFoundError(f"Input directory not found: {input_dir}")

        os.makedirs(output_dir, exist_ok=True)
        converted_files: Dict[str, str] = {}

        for filename in os.listdir(input_dir):
            if any(filename.lower().endswith(ext) for ext in file_extensions):
                input_filepath: str = os.path.join(input_dir, filename)
                output_filename: str = os.path.splitext(filename)[0] + ".md"
                output_filepath: str = os.path.join(output_dir, output_filename)

                if filename.lower().endswith(".pdf"):
                    _convert_pdf_to_markdown(input_filepath, output_filepath)
                elif filename.lower().endswith(".docx"):
                    _convert_docx_to_markdown(input_filepath, output_filepath)
                converted_files[input_filepath] = output_filepath

        return converted_files

    except FileNotFoundError as e:
        logger.error(e)
        raise
    except OSError as e:
        logger.error(f"OS error during file processing: {e}")
        raise
    except Exception as e:
        logger.exception("An unexpected error occurred:") #log full stack trace
        raise

def build_document_index(input_dir: str, output_dir: str) -> Tuple[List, Dict]:
    """ Asynchronously builds an index of Markdown documents from specified file types."""
    try:
        converted_docs: Dict[str, str] = convert_files_to_markdown(input_dir, output_dir, FILE_TYPES)
        documents: List[Document] = []
        term_frequency: Dict[str,int|Any] = {}


        for filename in os.listdir(output_dir):
            if filename.lower().endswith(".md") and not filename.lower().endswith("_toc.md"):
                filepath: str = os.path.join(output_dir, filename)
                content: str =  read_file_content(filepath)
                sections = split_content_into_sections(content)
                for section_title, section_content in sections:
                    metadata: Dict[str, str] = {"filename": filename, "section": section_title}
                    chunks: List[Document] = split_section_into_chunks((section_title, section_content), metadata)

                    for chunk in chunks:
                        chunk.update_derived_attributes()
                        unique_terms: Set[str] = {*chunk.clean_terms}
                        for term in unique_terms:
                            term_frequency[term] = term_frequency.get(term, 0) + 1
                    documents.extend(chunks)
        return documents, term_frequency
    except FileNotFoundError as e:
        logger.error(f"Input or output directory not found: {e}")
        raise
    except Exception:
        logger.exception("An unexpected error occurred during document index building:")
        raise


def split_section_into_chunks(section: Tuple[str, str], metadata: Dict[str, str]) -> List[Document]:
    """Splits a section into chunks of sentences."""
    section_title, section_content = section
    _sentences: List[str] = tokenize_sentences(section_content)
    _chunks: List[Document] = []
    _current_chunk: List[str] = []
    _current_length: int = 0

    try:
        for sentence in _sentences:
            _sentence_length: int = len(sentence)
            if _current_length + _sentence_length > chunk_size:
                chunk_content = f"Document: {metadata["filename"]}\nSection: {section_title}\n Snippet: {''.join(_current_chunk)}"
                _chunks.append(create_document_chunk(chunk_content, metadata))
                _current_chunk.clear()
                _current_chunk.append(sentence)
            else:
                _current_chunk.clear()
                _current_chunk.append(sentence)
                _current_length+=_sentence_length

        if _current_chunk:
            chunk_content = f"Document: {metadata["filename"]}\nSection: {section_title}\n Snippet: {''.join(_current_chunk)}"
            _chunks.append(create_document_chunk(chunk_content, metadata))

        return _chunks
    except FileNotFoundError as e:
        logger.error(f"Input or output directory not found: {e}")
        raise
    except Exception as e:
        logger.exception("An unexpected error occurred during document index building:")
        raise

def create_document_chunk(chunk_content: str, metadata: Dict[str, str]) -> Document:
    """Creates a Document object from chunk content and metadata."""
    return Document(chunk_content, metadata)


def tokenize_sentences(text: str) -> List[str]:
    """Tokenizes text into sentences."""
    return nltk.sent_tokenize(text, "english")


def split_content_into_sections(content: str) -> List[Tuple[str, str]]:
    """Splits content into sections based on Markdown headers (##, ***, etc.)."""
    sections: List[Tuple[str, str]] = []
    current_section: List[str] = []
    section_title: Optional[str] = None

    for line in content.splitlines():
        match = re.match(r"^([#|*|$]+)\s*(.*)", line) # Added \s* to handle whitespace
        if match and len(match.group(1)) > 1 and len(match.group(2).strip()) >= 5:
            if current_section and section_title is not None:
                sections.append((section_title, "\n".join(current_section[1:])))
                current_section = []
            section_title = match.group(2).strip()
        current_section.append(line)

    if current_section and section_title is not None:
        sections.append((section_title, "\n".join(current_section[1:])))
    return sections


def read_file_content(filepath: str) -> str:
    """ Reads the content of a file. """
    try:
        with open(filepath, 'r', encoding="UTF-8") as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error reading {filepath}: {e}")
        raise